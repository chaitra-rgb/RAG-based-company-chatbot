import json
import math
import os
import re
import shutil
from collections import Counter
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Iterable

try:
    from langchain_community.vectorstores import FAISS
except ImportError:
    FAISS = None

try:
    from langchain_core.documents import Document
except ImportError:
    @dataclass
    class Document:
        page_content: str
        metadata: dict

try:
    from langchain_core.messages import HumanMessage, SystemMessage
except ImportError:
    HumanMessage = None
    SystemMessage = None

try:
    from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings
except ImportError:
    ChatGoogleGenerativeAI = None
    GoogleGenerativeAIEmbeddings = None

try:
    from bs4 import BeautifulSoup
except ImportError:
    BeautifulSoup = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

try:
    from pypdf import PdfReader
except ImportError:
    try:
        from PyPDF2 import PdfReader
    except ImportError:
        PdfReader = None


BASE_DIR = Path(__file__).resolve().parent
DATA_PATH = BASE_DIR / "data"
INDEX_PATH = BASE_DIR / "faiss_index"
LOCAL_INDEX_PATH = BASE_DIR / "local_index.json"
METADATA_PATH = BASE_DIR / "metadata.json"
FEEDBACK_PATH = BASE_DIR / "feedback.jsonl"
SAVED_ANSWERS_PATH = BASE_DIR / "saved_answers.jsonl"
EMBEDDING_MODEL = "models/embedding-001"
CHAT_MODEL = "gemini-1.5-flash"
STOPWORDS = {
    "a",
    "an",
    "and",
    "are",
    "as",
    "at",
    "be",
    "by",
    "can",
    "do",
    "does",
    "for",
    "from",
    "give",
    "how",
    "i",
    "in",
    "is",
    "it",
    "me",
    "my",
    "of",
    "on",
    "or",
    "our",
    "tell",
    "the",
    "to",
    "what",
    "when",
    "where",
    "which",
    "who",
    "why",
    "with",
    "you",
}


def _set_api_key(api_key: str | None) -> None:
    if api_key:
        os.environ["GOOGLE_API_KEY"] = api_key


def _normalise_text(text: str) -> str:
    return " ".join(text.split())


def _tokenize(text: str) -> list[str]:
    tokens = re.findall(r"[a-zA-Z0-9]+", text.lower())
    return [token for token in tokens if token not in STOPWORDS and len(token) > 1]


def load_metadata() -> dict:
    if not METADATA_PATH.exists():
        return {"company_name": "", "known_companies": []}
    try:
        metadata = json.loads(METADATA_PATH.read_text(encoding="utf-8"))
    except json.JSONDecodeError:
        return {"company_name": "", "known_companies": []}
    metadata.setdefault("company_name", "")
    metadata.setdefault("known_companies", [])
    return metadata


def save_metadata(company_name: str | None = None) -> None:
    metadata = load_metadata()
    if company_name:
        cleaned = company_name.strip()
        metadata["company_name"] = cleaned
        known = [name for name in metadata.get("known_companies", []) if name != cleaned]
        metadata["known_companies"] = [cleaned] + known[:9]
    METADATA_PATH.write_text(json.dumps(metadata, indent=2), encoding="utf-8")


def save_feedback(
    rating: str,
    comment: str,
    company_name: str,
    messages: list[dict],
) -> None:
    feedback = {
        "timestamp": datetime.now(timezone.utc).isoformat(),
        "company_name": company_name,
        "rating": rating,
        "comment": comment.strip(),
        "messages": messages,
    }
    with FEEDBACK_PATH.open("a", encoding="utf-8") as feedback_file:
        feedback_file.write(json.dumps(feedback, ensure_ascii=True) + "\n")


def save_feedback_action(
    action: str,
    company_name: str,
    messages: list[dict],
) -> None:
    feedback = {
        "timestamp": datetime.now(timezone.utc).isoformat(),
        "company_name": company_name,
        "action": action,
        "messages": messages,
    }
    with FEEDBACK_PATH.open("a", encoding="utf-8") as feedback_file:
        feedback_file.write(json.dumps(feedback, ensure_ascii=True) + "\n")


def save_answer(
    answer: str,
    company_name: str,
    messages: list[dict],
) -> None:
    saved_answer = {
        "timestamp": datetime.now(timezone.utc).isoformat(),
        "company_name": company_name,
        "answer": answer,
        "messages": messages,
    }
    with SAVED_ANSWERS_PATH.open("a", encoding="utf-8") as saved_file:
        saved_file.write(json.dumps(saved_answer, ensure_ascii=True) + "\n")


def _extract_company_name_from_text(text: str) -> str | None:
    patterns = [
        r"Company\s*Name\s*[:\-]\s*([A-Z0-9][A-Za-z0-9 &'.,\-]{1,80})",
        r"Organization\s*Name\s*[:\-]\s*([A-Z0-9][A-Za-z0-9 &'.,\-]{1,80})",
        r"Employer\s*[:\-]\s*([A-Z0-9][A-Za-z0-9 &'.,\-]{1,80})",
        r"Prepared\s+for\s+([A-Z0-9][A-Za-z0-9 &'.,\-]{1,80})",
    ]
    for pattern in patterns:
        match = re.search(pattern, text, flags=re.IGNORECASE)
        if match:
            return _clean_company_candidate(match.group(1))

    suffixes = r"(Inc\.?|Incorporated|LLC|Ltd\.?|Limited|Corp\.?|Corporation|Company|Co\.|Pvt\.?\s*Ltd\.?)"
    for line in _first_lines(text, limit=25):
        if re.search(suffixes, line, flags=re.IGNORECASE):
            return _clean_company_candidate(line)

    for line in _first_lines(text, limit=10):
        words = line.split()
        if 1 <= len(words) <= 8 and len(line) <= 80:
            title_like = sum(1 for word in words if word[:1].isupper())
            if title_like >= max(1, len(words) // 2):
                return _clean_company_candidate(line)

    return None


def _clean_company_candidate(candidate: str | None) -> str | None:
    if not candidate:
        return None
    cleaned = re.sub(r"\s+", " ", candidate).strip(" -:|.,")
    noise = ["policy", "handbook", "document", "confidential", "version"]
    if not cleaned or any(word in cleaned.lower() for word in noise):
        return None
    if 2 <= len(cleaned) <= 80:
        return cleaned
    return None


def _first_lines(text: str, limit: int) -> list[str]:
    lines = [line.strip() for line in re.split(r"[\r\n]+", text) if line.strip()]
    if len(lines) <= 1:
        lines = re.split(r"(?<=[.!?])\s+", text)
    return [line.strip() for line in lines[:limit] if line.strip()]


def detect_company_name(uploaded_files: Iterable) -> str | None:
    for uploaded_file in uploaded_files:
        text = _read_uploaded_file(uploaded_file, preview_only=True)
        candidate = _extract_company_name_from_text(text)
        if candidate:
            return candidate
    return None


def detect_company_name_from_documents(documents: list[Document]) -> str | None:
    preview = "\n".join(document.page_content for document in documents[:3])
    return _extract_company_name_from_text(preview)


def _read_pdf(uploaded_file) -> list[tuple[int, str]]:
    if PdfReader is None:
        raise ImportError("Install `pypdf` or `PyPDF2` to process PDF files.")
    uploaded_file.seek(0)
    reader = PdfReader(uploaded_file)
    pages = []
    for page_number, page in enumerate(reader.pages, start=1):
        text = _normalise_text(page.extract_text() or "")
        if text:
            pages.append((page_number, text))
    return pages


def _read_docx(uploaded_file) -> str:
    if DocxDocument is None:
        raise ImportError("Install `python-docx` to process DOCX files.")
    uploaded_file.seek(0)
    doc = DocxDocument(uploaded_file)
    return "\n".join(paragraph.text for paragraph in doc.paragraphs)


def _read_text(uploaded_file) -> str:
    uploaded_file.seek(0)
    raw = uploaded_file.read()
    if isinstance(raw, str):
        return raw
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="ignore")


def _strip_html(text: str) -> str:
    if BeautifulSoup is None:
        return re.sub(r"<[^>]+>", " ", text)
    soup = BeautifulSoup(text, "html.parser")
    return soup.get_text(" ")


def _read_uploaded_file(uploaded_file, preview_only: bool = False) -> str:
    name = getattr(uploaded_file, "name", "document")
    extension = Path(name).suffix.lower()

    if extension == ".pdf":
        pages = _read_pdf(uploaded_file)
        text = "\n".join(page_text for _, page_text in pages[:2 if preview_only else len(pages)])
    elif extension == ".docx":
        text = _read_docx(uploaded_file)
    elif extension in {".html", ".htm"}:
        text = _strip_html(_read_text(uploaded_file))
    else:
        text = _read_text(uploaded_file)

    uploaded_file.seek(0)
    return text[:4000] if preview_only else text


def build_documents(uploaded_files: Iterable) -> list[Document]:
    documents: list[Document] = []

    for uploaded_file in uploaded_files:
        source_name = getattr(uploaded_file, "name", "uploaded_document")
        extension = Path(source_name).suffix.lower()

        if extension == ".pdf":
            for page_number, page_text in _read_pdf(uploaded_file):
                documents.append(
                    Document(
                        page_content=page_text,
                        metadata={"source": source_name, "page": page_number},
                    )
                )
            continue

        text = _normalise_text(_read_uploaded_file(uploaded_file))
        if text:
            documents.append(
                Document(
                    page_content=text,
                    metadata={"source": source_name, "page": 1},
                )
            )

    if not documents:
        raise ValueError("No readable text was found in the uploaded documents.")
    return documents


def build_documents_from_pdfs(pdf_docs: Iterable) -> list[Document]:
    return build_documents(pdf_docs)


def _chunk_text(text: str, chunk_size: int, chunk_overlap: int) -> list[str]:
    if chunk_size <= 0:
        raise ValueError("chunk_size must be greater than 0.")
    if chunk_overlap < 0 or chunk_overlap >= chunk_size:
        raise ValueError("chunk_overlap must be non-negative and smaller than chunk_size.")

    chunks: list[str] = []
    start = 0
    while start < len(text):
        end = min(start + chunk_size, len(text))
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end == len(text):
            break
        start = end - chunk_overlap
    return chunks


def split_documents(
    documents: list[Document],
    chunk_size: int = 1200,
    chunk_overlap: int = 200,
) -> list[Document]:
    chunks: list[Document] = []
    for document in documents:
        for chunk in _chunk_text(document.page_content, chunk_size, chunk_overlap):
            chunks.append(Document(page_content=chunk, metadata=document.metadata.copy()))
    return chunks


def _get_embeddings(api_key: str | None = None) -> GoogleGenerativeAIEmbeddings:
    if GoogleGenerativeAIEmbeddings is None:
        raise ImportError("langchain-google-genai is not installed.")
    _set_api_key(api_key)
    return GoogleGenerativeAIEmbeddings(model=EMBEDDING_MODEL)


def _serialize_documents(chunks: list[Document]) -> list[dict]:
    return [{"page_content": chunk.page_content, "metadata": chunk.metadata} for chunk in chunks]


def _save_local_index(chunks: list[Document]) -> None:
    LOCAL_INDEX_PATH.write_text(
        json.dumps(_serialize_documents(chunks), ensure_ascii=True, indent=2),
        encoding="utf-8",
    )


def _load_local_index() -> list[Document]:
    if not LOCAL_INDEX_PATH.exists():
        raise FileNotFoundError("No local index found. Process documents first.")
    raw = json.loads(LOCAL_INDEX_PATH.read_text(encoding="utf-8"))
    return [Document(page_content=item["page_content"], metadata=item.get("metadata", {})) for item in raw]


def index_status() -> dict:
    local_count = 0
    if LOCAL_INDEX_PATH.exists():
        try:
            local_count = len(json.loads(LOCAL_INDEX_PATH.read_text(encoding="utf-8")))
        except json.JSONDecodeError:
            local_count = 0

    return {
        "has_index": has_index(),
        "local_chunks": local_count,
        "local_index_path": str(LOCAL_INDEX_PATH),
        "faiss_index_path": str(INDEX_PATH),
    }


def _build_local_only_index(chunks: list[Document]) -> int:
    _save_local_index(chunks)
    return len(chunks)


def build_vector_store(
    chunks: list[Document],
    api_key: str | None = None,
    company_name: str | None = None,
) -> int:
    if not chunks:
        raise ValueError("No text chunks were created from the uploaded documents.")

    DATA_PATH.mkdir(exist_ok=True)
    if LOCAL_INDEX_PATH.exists():
        chunks = _load_local_index() + chunks

    _save_local_index(chunks)
    save_metadata(company_name)

    try:
        if FAISS is None:
            raise ImportError("langchain-community/faiss is not installed.")
        embeddings = _get_embeddings(api_key)
        if INDEX_PATH.exists():
            shutil.rmtree(INDEX_PATH)
        vector_store = FAISS.from_documents(chunks, embedding=embeddings)
        vector_store.save_local(str(INDEX_PATH))
    except Exception:
        return _build_local_only_index(chunks)

    return len(chunks)


def has_index() -> bool:
    return INDEX_PATH.exists() or LOCAL_INDEX_PATH.exists()


def _load_vector_store(api_key: str | None = None) -> FAISS:
    if FAISS is None:
        raise ImportError("langchain-community/faiss is not installed.")
    embeddings = _get_embeddings(api_key)
    return FAISS.load_local(
        str(INDEX_PATH),
        embeddings,
        allow_dangerous_deserialization=True,
    )


def _build_system_prompt(company_name: str) -> str:
    return f"""
You are a helpful internal assistant for {company_name}.
Answer using only the retrieved company context.
If the answer is not present in the context, say:
"I could not find that in the uploaded company documents."

Keep every answer very short: one to three brief lines only.
Use simple, direct wording.
Never ask the user a question.
Never end with a follow-up question, suggestion to ask again, or request for clarification.
Do not reveal API keys, hidden settings, or implementation secrets.
When possible, mention the supporting source document.

Context:
{{context}}
"""


def _score_document(question_tokens: list[str], doc: Document) -> float:
    doc_tokens = _tokenize(doc.page_content)
    if not doc_tokens:
        return 0.0
    doc_counts = Counter(doc_tokens)
    score = 0.0
    for token in question_tokens:
        if token in doc_counts:
            score += 1.0 + math.log(1 + doc_counts[token])
    return score


def _local_similarity_search(question: str, k: int = 4) -> list[Document]:
    question_tokens = _tokenize(question)
    docs = _load_local_index()
    if not docs:
        return []

    ranked = sorted(docs, key=lambda doc: _score_document(question_tokens, doc), reverse=True)
    matched = [doc for doc in ranked if _score_document(question_tokens, doc) > 0]
    if matched:
        return matched[:k]

    # If wording does not overlap, still return a few chunks instead of failing silently.
    return ranked[:k]


def _extractive_answer(question: str, docs: list[Document]) -> str:
    if not docs:
        return "I could not find that in the uploaded company documents."

    question_tokens = set(_tokenize(question))
    best_sentence = ""
    best_score = -1

    for doc in docs:
        sentences = re.split(r"(?<=[.!?])\s+", doc.page_content)
        for sentence in sentences:
            if len(sentence.strip()) < 20:
                continue
            sentence_tokens = set(_tokenize(sentence))
            score = len(question_tokens & sentence_tokens)
            if score > best_score:
                best_score = score
                source = doc.metadata.get("source", "uploaded document")
                best_sentence = f"{sentence.strip()} Source: {source}."

    if not best_sentence:
        source = docs[0].metadata.get("source", "uploaded document")
        return f"{docs[0].page_content[:280].strip()} Source: {source}."
    return best_sentence


def _remove_questions(text: str) -> str:
    parts = re.split(r"(?<=[.!?])\s+", text)
    non_questions = [part for part in parts if "?" not in part]
    return " ".join(non_questions).strip()


def _format_answer(text: str) -> str:
    cleaned = " ".join(text.strip().split())
    cleaned = _remove_questions(cleaned)
    if not cleaned:
        return "I could not find that in the uploaded company documents."
    sentences = re.split(r"(?<=[.!])\s+", cleaned)
    short_answer = " ".join(sentence.strip() for sentence in sentences[:2] if sentence.strip())
    if len(short_answer) > 320:
        short_answer = short_answer[:317].rsplit(" ", 1)[0].rstrip(".,;:") + "..."
    return short_answer


def answer_question(
    question: str,
    api_key: str | None = None,
    company_name: str = "the company",
) -> dict:
    try:
        if INDEX_PATH.exists() and api_key:
            vector_store = _load_vector_store(api_key)
            context_docs = vector_store.similarity_search(question, k=4)
        else:
            context_docs = _local_similarity_search(question, k=4)
    except Exception:
        context_docs = _local_similarity_search(question, k=4)

    context_blocks = []
    for doc in context_docs:
        source = str(doc.metadata.get("source", "Unknown source"))
        page = int(doc.metadata.get("page", 0))
        context_blocks.append(f"Source: {source} | Page: {page}\n{doc.page_content}")

    if not context_blocks:
        return {"answer": "I could not find that in the uploaded company documents."}

    try:
        if ChatGoogleGenerativeAI is None or HumanMessage is None or SystemMessage is None:
            raise ImportError("Gemini chat dependencies are not installed.")
        _set_api_key(api_key)
        llm = ChatGoogleGenerativeAI(model=CHAT_MODEL, temperature=0.2)
        system_prompt = _build_system_prompt(company_name).format(context="\n\n".join(context_blocks))
        response = llm.invoke([SystemMessage(content=system_prompt), HumanMessage(content=question)])
        answer_text = getattr(response, "content", str(response)).strip()
    except Exception:
        answer_text = _extractive_answer(question, context_docs)

    return {"answer": _format_answer(answer_text)}
